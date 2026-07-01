from __future__ import annotations

import io
import logging
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import xlrd
from docx import Document
from openpyxl import load_workbook

from src.config import TZ_OCR_LANG, TZ_OCR_PSM, TZ_OCR_SCALE, TZ_PDF_OCR_ENABLED, TESSERACT_CMD
from src.services.models import TZItem

logger = logging.getLogger(__name__)

SUPPORTED_TZ_EXTENSIONS = {".doc", ".docx", ".pdf", ".xlsx", ".xls"}
SUPPORTED_TZ_LABEL = ".doc, .docx, .pdf, .xlsx, .xls"

_PRODUCT_UNITS = frozenset(
    {
        "шт",
        "шт.",
        "штука",
        "штук",
        "компл",
        "компл.",
        "комплект",
        "упак",
        "упак.",
        "упаковка",
        "м",
        "м.",
        "метр",
        "кг",
        "кг.",
        "л",
        "л.",
        "т",
        "т.",
        "пара",
        "пар",
        "набор",
    }
)


@dataclass
class _ColumnMap:
    number: int
    name: int
    unit: int
    qty: int
    specs: int | None = None
    country: int | None = None
    sale_price: int | None = None
    has_number: bool = True


def parse_tz(path: Path) -> list[TZItem]:
    suffix = path.suffix.lower()
    if suffix == ".doc":
        return _parse_tz_doc(path)
    if suffix == ".docx":
        return _parse_tz_docx(path)
    if suffix == ".pdf":
        return _parse_tz_pdf(path)
    if suffix in {".xlsx", ".xls"}:
        return _parse_tz_excel(path)
    raise ValueError(
        f"Неподдерживаемый формат ТЗ: {suffix}. "
        f"Допустимые форматы: {SUPPORTED_TZ_LABEL}"
    )


def extract_tz_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".doc":
        return _extract_doc_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".pdf":
        return _extract_pdf_document_text(path)
    if suffix in {".xlsx", ".xls"}:
        return _extract_excel_text(path)
    raise ValueError(
        f"Неподдерживаемый формат ТЗ: {suffix}. "
        f"Допустимые форматы: {SUPPORTED_TZ_LABEL}"
    )


def detect_tz_suffix(content: bytes, content_type: str | None = None) -> str | None:
    if content.startswith(b"%PDF"):
        return ".pdf"
    if content.startswith(b"\xd0\xcf\x11\xe0"):
        return _detect_ole_suffix(content)
    if content.startswith(b"PK"):
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = archive.namelist()
                if any(name.startswith("word/") for name in names):
                    return ".docx"
                if any(name.startswith("xl/") for name in names):
                    return ".xlsx"
        except zipfile.BadZipFile:
            pass

    if content_type:
        lowered = content_type.lower()
        if "pdf" in lowered:
            return ".pdf"
        if "spreadsheet" in lowered or "excel" in lowered:
            return ".xlsx"
        if "word" in lowered or "document" in lowered:
            return ".docx"

    return None


def resolve_tz_upload_filename(
    filename: str | None,
    content: bytes,
    content_type: str | None = None,
) -> str:
    name = (filename or "tz").strip() or "tz"
    suffix = Path(name).suffix.lower()
    if suffix in SUPPORTED_TZ_EXTENSIONS:
        return name

    detected = detect_tz_suffix(content, content_type)
    if detected:
        stem = Path(name).stem or "tz"
        return f"{stem}{detected}"

    raise ValueError(f"Загрузите файл ТЗ: {SUPPORTED_TZ_LABEL}")


def _cell_str(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


_HEADER_NAME_TOKENS = frozenset(
    {
        "наименование",
        "товара",
        "характеристики",
        "характеристика",
        "характеристи",
        "характери",
        "единиц",
        "единица",
        "измерения",
        "измерение",
        "кол-во",
        "количество",
        "страны",
        "страна",
        "происхождения",
        "значение",
        "п/п",
        "итого",
        "стики",
        "ние",
    }
)


def _normalize_tz_product_name(name: str) -> str:
    cleaned = name.strip().strip("|").strip()
    cleaned = re.sub(r"^\d+\s*\|", "", cleaned).strip()
    cleaned = cleaned.rstrip("|").strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned


def _is_valid_tz_product_name(name: str) -> bool:
    cleaned = _normalize_tz_product_name(name)
    if len(cleaned) < 4:
        return False

    lower = cleaned.lower()
    if lower in _HEADER_NAME_TOKENS:
        return False

    words = [word for word in re.split(r"[\s|]+", lower) if word]
    if words and all(word in _HEADER_NAME_TOKENS for word in words):
        return False
    if len(words) == 1:
        if words[0] in _HEADER_NAME_TOKENS or len(cleaned) < 5:
            return False
    elif len(words) < 2 and len(cleaned) < 12:
        return False
    if re.search(r"(?:федерация|российск)", lower):
        return False
    if lower in {"мпиксель", "крат", "прямой", "верхнее", "нижнее", "монокулярный", "светодиод", "cmos", "usb"}:
        return False
    if cleaned.count("|") >= 2 and len(cleaned) < 40:
        return False
    if not re.search(r"[а-яa-zё]{3,}", lower):
        return False
    if any(
        token in lower
        for token in (
            "госпрограмма",
            "зарница",
            "купить учебное",
            "основной перечень",
        )
    ):
        return False

    return True


def _is_tz_table_header(header_cells: list[str]) -> bool:
    joined = " ".join(cell.strip().lower() for cell in header_cells if cell.strip())
    if "наимен" not in joined and not _looks_like_price_request_header(joined):
        return False
    if _looks_like_price_request_header(joined):
        return _build_price_request_column_map(header_cells) is not None
    return _build_column_map(header_cells) is not None


_VENDOR_CODE_RE = re.compile(r"\b([A-Z]{2,5}\d{3,7})\b")
_PRICE_REQUEST_MARKERS = (
    "коммерческ",
    "koммерч",
    "koммep",
    "ценов",
    "wehoby",
    "wehoв",
    "ценовой запрос",
    "предоставлен",
    "npedocma",
)


def _looks_like_price_request_header(header: str) -> bool:
    lower = header.lower()
    if "товар" in lower:
        return False
    if ("ед" in lower and "изм" in lower) or "единиц" in lower:
        return False
    has_name = "наимен" in lower or "haumen" in lower or "haumeh" in lower
    has_price = any(marker in lower for marker in ("цен", "lena", "сумм", "cymma"))
    if not has_name or not has_price:
        return False
    return True


def _looks_like_price_request_document(text: str) -> bool:
    if len(_VENDOR_CODE_RE.findall(text)) < 1:
        return False
    lower = text.lower()
    if re.search(r"(?:^|[\n\r])\s*\d+\s*\|\s*[A-Z]{2,5}\d{3,7}\b", text, re.MULTILINE):
        return True
    if re.search(r"(?:^|[\n\r])\s*\d+\s+[A-Z]{2,5}\d{3,7}\b", text, re.MULTILINE):
        return any(marker in lower for marker in _PRICE_REQUEST_MARKERS)
    return False


@dataclass
class _PriceRequestColumnMap:
    number: int
    name: int
    qty: int | None = None


def _build_price_request_column_map(header_cells: list[str]) -> _PriceRequestColumnMap | None:
    lower = [cell.strip().lower() for cell in header_cells]
    name_idx: int | None = None
    for index, cell in enumerate(lower):
        if "наимен" in cell or "haumen" in cell or "haumeh" in cell:
            name_idx = index
            break
    if name_idx is None:
        return None

    number_idx = 0
    for index, cell in enumerate(lower):
        if cell in {"№", "п/п", "пп", "n", "no", "wn"} or cell.startswith("№"):
            number_idx = index
            break

    qty_idx: int | None = None
    for index, cell in enumerate(lower):
        if "кол" in cell or "kon" in cell:
            qty_idx = index
            break

    return _PriceRequestColumnMap(number=number_idx, name=name_idx, qty=qty_idx)


_OCR_LATIN_TO_CYRILLIC = str.maketrans(
    {
        "A": "А",
        "B": "В",
        "C": "С",
        "E": "Е",
        "H": "Н",
        "K": "К",
        "M": "М",
        "O": "О",
        "P": "Р",
        "T": "Т",
        "X": "Х",
        "Y": "У",
        "a": "а",
        "c": "с",
        "e": "е",
        "o": "о",
        "p": "р",
        "x": "х",
        "y": "у",
        "k": "к",
        "m": "м",
        "t": "т",
        "s": "с",
        "u": "у",
        "i": "и",
        "l": "л",
        "n": "н",
        "d": "д",
        "g": "г",
        "r": "р",
        "f": "ф",
        "v": "в",
        "w": "в",
        "z": "з",
        "j": "й",
        "q": "я",
        "h": "н",
        "b": "в",
    }
)

_OCR_CYRILLIC_TO_LATIN = str.maketrans(
    {
        "А": "A",
        "В": "B",
        "С": "C",
        "Е": "E",
        "Н": "H",
        "К": "K",
        "М": "M",
        "О": "O",
        "Р": "P",
        "Т": "T",
        "Х": "X",
        "У": "Y",
        "а": "a",
        "в": "b",
        "с": "c",
        "е": "e",
        "н": "h",
        "к": "k",
        "м": "m",
        "о": "o",
        "р": "p",
        "т": "t",
        "х": "x",
        "у": "y",
        "и": "i",
        "л": "l",
        "д": "d",
        "г": "g",
        "ф": "f",
        "з": "z",
        "й": "j",
        "я": "q",
        "ш": "w",
        "щ": "w",
        "ь": "",
        "ъ": "",
        "ы": "y",
        "э": "e",
        "ю": "u",
        "ё": "e",
    }
)

_KIT_ELEMENT_MARKERS = (
    r"элемент|"
    r"3[l1i][eе][mм][eе][nн][tт]|"
    r"3ЛЕМ|"
    r"3IIЕМ|"
    r"3IИЕМ|"
    r"9II@|"
    r"ЛЕМЕН|"
    r"J[LlЛ1][EeЕе][MmМm][EeЕе][NnНn]|"
    r"JЛЕМ|"
    r"@MCH|"
    r"@МСН|"
    r"МСНТ|"
    r"9JIEM|"
    r"9II@MCH|"
    r"element|"
    r"EHTOV|EHTOB|EMEH|@MCH|"
    r"карто"
)

_OCR_WORD_FIXES: dict[str, str] = {
    "po6or": "робот",
    "po6ot": "робот",
    "po6o": "робо",
    "borm": "ботли",
    "botli": "ботли",
    "berm": "ботли",
    "bota": "ботли",
    "bora": "бота",
    "paspuparomaa": "развивающая",
    "paspuparomat": "развивающая",
    "passaparomaa": "развивающая",
    "passaparomat": "развивающая",
    "raspuraromaa": "развивающая",
    "raspuraromat": "развивающая",
    "urpyuma": "игрушка",
    "urpyuika": "игрушка",
    "urpyuka": "игрушка",
    "urpyua": "игрушка",
    "urruuma": "игрушка",
    "yrryuma": "игрушка",
    "axceccyapbi": "аксессуары",
    "akceccyapbi": "аксессуары",
    "axceccyapsi": "аксессуары",
    "axceccyapbl": "аксессуары",
    "axcessuarybi": "аксессуары",
    "axceccuarybi": "аксессуары",
    "crpountes": "строитель",
    "crpoutems": "строитель",
    "crpoums": "строитель",
    "crpoumes": "строитель",
    "crpoumems": "строитель",
    "kommiexr": "комплект",
    "kommexr": "комплект",
    "kommaext": "комплект",
    "kommiext": "комплект",
    "kommieht": "комплект",
    "tematuyeckux": "тематических",
    "tematuyecskux": "тематических",
    "tematuyeckuxx": "тематических",
    "tematuyesckmx": "тематических",
    "tematuyesckmxx": "тематических",
    "temathyecckmxh": "тематических",
    "temathyecckmxx": "тематических",
    "tomel": "полей",
    "tome": "полей",
    "hoel": "полей",
    "hoel": "полей",
    "jlentoxc": "делюкс",
    "jlemoxc": "делюкс",
    "dentoxc": "делюкс",
    "jemoxc": "делюкс",
    "lenetoxc": "делюкс",
    "lentoxc": "делюкс",
    "лentoxc": "делюкс",
    "bepcua": "версия",
    "bepcusa": "версия",
    "bepcus": "версия",
    "bercua": "версия",
    "bercusa": "версия",
    "po6omsmubto": "робомышью",
    "po6omsuupt0": "робомышью",
    "po6ompmubio": "робомышью",
    "po6bompmubio": "робомышью",
    "po6boMpmubio": "робомышью",
    "ctpomm": "строим",
    "ctromm": "строим",
    "ctpoim": "строим",
    "ctpoMM": "строим",
    "maprlpytbi": "маршруты",
    "mapulpytbi": "маршруты",
    "maplipyTBI": "маршруты",
    "mapLIpyTBI": "маршруты",
    "maplipytbi": "маршруты",
    "duia": "для",
    "dua": "для",
    "quia": "для",
    "pobota": "робота",
    "pobora": "робота",
    "kaptoqkami": "карточками",
    "kapto ukami": "карточками",
    "delioks": "делюкс",
    "botley": "botley",
    "botli": "ботли",
    "vorm": "ботли",
}


def _ocr_word_fix_lookup(word: str) -> str:
    return word.translate(_OCR_CYRILLIC_TO_LATIN).lower().strip(".,;:\"«»-'")


def _apply_ocr_word_fix(word: str, *lookups: str) -> str | None:
    keys = [_ocr_word_fix_lookup(word), word.lower().strip(".,;:\"«»-'")]
    keys.extend(lookups)
    for key in keys:
        if not key:
            continue
        fixed = _OCR_WORD_FIXES.get(key)
        if fixed:
            if word[:1].isupper():
                return fixed[:1].upper() + fixed[1:]
            return fixed
    return None


def _repair_ocr_cyrillic_word(word: str) -> str:
    if not word:
        return word
    if re.search(r"<<<VC\d+>>>", word):
        return word
    if re.fullmatch(r"[A-Z]{2,5}\d{3,7}", word):
        return word
    if re.fullmatch(r"[\d.,]+", word):
        return word

    repaired = re.sub(r"(?<=[A-Za-zА-Яа-я])6(?=[A-Za-zА-Яа-я])", "б", word)
    repaired = re.sub(r"J[lI1]", "Л", repaired)
    repaired = re.sub(r"(?<=[A-Za-zА-Яa-я])r(?=$)", "т", repaired)
    repaired = repaired.translate(_OCR_LATIN_TO_CYRILLIC)
    repaired = re.sub(r"J[lI1]", "Л", repaired)

    fixed = _apply_ocr_word_fix(word, _ocr_word_fix_lookup(word))
    if fixed:
        return fixed
    fixed = _apply_ocr_word_fix(repaired, _ocr_word_fix_lookup(repaired))
    if fixed:
        return fixed

    if re.search(r"[A-Za-z]", repaired) and re.search(r"[А-Яа-я]", repaired):
        merged = repaired.translate(_OCR_LATIN_TO_CYRILLIC)
        fixed = _apply_ocr_word_fix(merged, _ocr_word_fix_lookup(merged))
        if fixed:
            return fixed
        return merged
    return repaired


def _repair_ocr_cyrillic_text(text: str) -> str:
    if not text.strip():
        return text

    protected: dict[str, str] = {}

    def _protect(match: re.Match[str]) -> str:
        token = f"<<<VC{len(protected)}>>>"
        protected[token] = match.group(0)
        return token

    shielded = re.sub(r"\b[A-Z]{2,5}\d{3,7}\b", _protect, text)
    words = re.split(r"(\s+)", shielded)
    repaired_words = [
        _repair_ocr_cyrillic_word(part) if part.strip() and not part.isspace() else part
        for part in words
    ]
    repaired = "".join(repaired_words)
    for token, original in protected.items():
        repaired = repaired.replace(token, original)
    return repaired


def _trim_price_request_name_segment(segment: str) -> str:
    trimmed = segment.strip()
    kit_match = re.search(
        rf"\(\s*\d+\s*[\s\S]{{0,80}}?(?:{_KIT_ELEMENT_MARKERS})",
        trimmed,
        re.I,
    )
    if kit_match:
        trimmed = trimmed[: kit_match.start()].strip()
    return trimmed


def _extract_quoted_product_name(segment: str) -> str:
    quoted_parts: list[str] = []
    for pattern in (
        r'"([^"]{2,500})"',
        r"'([^']{2,500})'",
        r"«([^»]{2,500})»",
    ):
        for match in re.finditer(pattern, segment, re.S):
            cleaned = re.sub(r"\s+", " ", match.group(1)).strip(" .")
            if cleaned and cleaned not in quoted_parts:
                quoted_parts.append(cleaned)

    if not quoted_parts and '"' in segment:
        unclosed = re.search(
            rf'"([\s\S]{{2,500}}?)(?=\(\s*\d+\s*[\s\S]{{0,20}}?(?:{_KIT_ELEMENT_MARKERS})|\Z)',
            segment,
            re.I,
        )
        if unclosed:
            cleaned = re.sub(r"\s+", " ", unclosed.group(1)).strip(" .")
            if cleaned:
                quoted_parts.append(cleaned)

    return " ".join(quoted_parts)


_PRICE_REQUEST_PHRASE_FIXES: tuple[tuple[str, str], ...] = (
    (r"робота\s+бота\b", "робота Ботли"),
    (r"робота\s+bota\b", "робота Ботли"),
    (r"робота\s+bora\b", "робота Ботли"),
)


def _polish_price_request_name(name: str, vendor_code: str) -> str:
    if not name.upper().startswith(vendor_code.upper()):
        return name
    tail = name[len(vendor_code) :].strip()
    if not tail:
        return name
    polished_tail = " ".join(_repair_ocr_cyrillic_word(word) for word in tail.split())
    for pattern, replacement in _PRICE_REQUEST_PHRASE_FIXES:
        polished_tail = re.sub(pattern, replacement, polished_tail, flags=re.I)
    polished_tail = re.sub(r"\s+", " ", polished_tail).strip(' "')
    return f"{vendor_code} {polished_tail}".strip()


def _extract_price_request_name(segment: str, vendor_code: str) -> str:
    work = _trim_price_request_name_segment(segment)
    work = re.sub(re.escape(vendor_code), "", work, count=1, flags=re.I).strip()
    work = re.sub(r"^[|:\-\s]+", "", work)

    quoted = _extract_quoted_product_name(work)
    prefix = work
    first_quote = re.search(r'["«\']', work)
    if first_quote:
        prefix = work[: first_quote.start()]
    prefix = re.sub(r"\s+", " ", prefix).strip(' "|')

    parts = [part for part in (prefix, quoted) if part]
    if parts:
        combined = " ".join(parts)
        combined = re.sub(r"\s+", " ", combined).strip()
        return _polish_price_request_name(f"{vendor_code} {combined}".strip(), vendor_code)
    return vendor_code


def _parse_price_request_quantity(segment: str) -> float:
    """Кол-во заказа (шт.), не путать с «N элементов в комплекте»."""
    for pattern in (
        r"(?:кол[- ]?во|kon[- ]?bo)\s*[^\d]{0,12}(\d+(?:[.,]\d+)?)",
        r"\(\s*(\d+)\s*(?:шт\.?|ur\.?)\s*\)",
        r"(?:\|\s*){1,3}(\d+(?:[.,]\d+)?)\s*(?:шт\.?|ur\.?)?\s*(?:\||$)",
        r"(?<![(\d])(\d+(?:[.,]\d+)?)\s*(?:шт\.?|ur\.?)\s*$",
    ):
        match = re.search(pattern, segment.strip(), re.I | re.S)
        if match:
            try:
                value = float(match.group(1).replace(",", "."))
                if 0 < value <= 9999:
                    return value
            except ValueError:
                continue
    return 1.0


def _parse_price_request_kit_elements(segment: str) -> int | None:
    match = re.search(
        rf"\(\s*(\d+)\s*[\s\S]{{0,80}}?(?:{_KIT_ELEMENT_MARKERS})",
        segment,
        re.I,
    )
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def _parse_price_request_text(text: str) -> list[TZItem]:
    if not text.strip():
        return []

    text = _repair_ocr_cyrillic_text(text)

    item_pattern = re.compile(
        r"(?:^|[\n\r])\s*(\d+)\s*\|\s*([A-Z]{2,5}\d{3,7})\b",
        re.MULTILINE,
    )
    matches = list(item_pattern.finditer(text))
    if len(matches) < 1:
        item_pattern = re.compile(
            r"(?:^|[\n\r])\s*(\d+)\s+([A-Z]{2,5}\d{3,7})\b",
            re.MULTILINE,
        )
        matches = [match for match in item_pattern.finditer(text) if int(match.group(1)) <= 50]

    if not matches or not _looks_like_price_request_document(text):
        return []

    items: list[TZItem] = []
    for index, match in enumerate(matches):
        number = int(match.group(1))
        if number <= 0 or number > 200:
            continue
        vendor_code = match.group(2).upper()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        segment = text[start:end]
        itogo = re.search(r"\b(?:ИТОГО|uTOTO|ITOGO)\b", segment, re.I)
        if itogo:
            segment = segment[: itogo.start()]

        name = _extract_price_request_name(segment, vendor_code)
        if not _is_valid_tz_product_name(name) and name == vendor_code:
            continue

        specs_parts = [f"Код производителя: {vendor_code}"]
        kit_elements = _parse_price_request_kit_elements(segment)
        if kit_elements is not None:
            specs_parts.append(f"Комплектация: {kit_elements} элементов")

        items.append(
            TZItem(
                number=number,
                name=name,
                unit="шт.",
                quantity=_parse_price_request_quantity(segment),
                specifications="; ".join(specs_parts),
            )
        )

    return items


_NUMBERED_LIST_ITEM_RE = re.compile(
    r"^\s*(?:(?P<dup>\d+)\s+)?(?P<num>\d+)\s*[\.\)]\s*(?:\.\s*)?(?P<name>.+?)\.?\s*$"
)
_NUMBERED_LIST_SKIP_MARKERS = (
    "приложение",
    "описание объекта",
    "форма коммерческ",
    "форма коммерч",
    "итого",
    "наименование",
    "№ п/п",
    "кол-во",
    "количество",
)
_NUMBERED_LIST_SECTION_HEADERS = (
    "базовый (практический) комплект",
    "базовый комплект",
    "практический комплект",
)


def _looks_like_numbered_equipment_list(text: str) -> bool:
    if not text.strip():
        return False
    lower = text.lower()
    if any(marker in lower for marker in ("наименование", "№ п/п", "кол-во", "количество")):
        if "наимен" in lower and ("кол" in lower or "ед" in lower):
            return False
    entries = _collect_numbered_list_entries(text)
    numbered_valid = [
        entry
        for entry in entries
        if entry[0] is not None and _is_valid_numbered_list_name(entry[1])
    ]
    return len(numbered_valid) >= 3


def _is_valid_numbered_list_name(name: str) -> bool:
    cleaned = _normalize_tz_product_name(name)
    if not _is_valid_tz_product_name(cleaned):
        return False
    lower = cleaned.lower()
    if any(marker in lower for marker in _NUMBERED_LIST_SKIP_MARKERS):
        return False
    if lower.rstrip(".") in _NUMBERED_LIST_SECTION_HEADERS:
        return False
    if re.search(r"@|mail\.ru|yandex|инн\s*:|егр[юл]", lower):
        return False
    if any(
        token in lower
        for token in (
            "министерство",
            "казенное",
            "учреждение",
            "област",
            "ответствен",
            "директор",
            "заместител",
            "тел.",
            "телефон",
            "e-mail",
            "email",
            "факс",
        )
    ):
        return False
    if re.search(r"\d{5,}", cleaned):
        return False
    if not any("а" <= char <= "я" or char == "ё" for char in lower):
        return False
    latin_letters = sum(1 for char in cleaned if char.isascii() and char.isalpha())
    if latin_letters / max(len(cleaned), 1) > 0.22:
        return False
    words = [word for word in re.split(r"[\s|]+", lower) if len(word) > 2]
    if len(words) < 2 and len(cleaned) < 22:
        return False
    return True


def _is_numbered_list_section_header(name: str) -> bool:
    lower = _normalize_tz_product_name(name).lower().rstrip(".")
    return lower in _NUMBERED_LIST_SECTION_HEADERS


def _collect_numbered_list_entries(text: str) -> list[tuple[int | None, str]]:
    entries: list[tuple[int | None, str]] = []
    current_num: int | None = None
    current_parts: list[str] = []
    numbered_seen = 0

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        match = _NUMBERED_LIST_ITEM_RE.match(line)
        if match:
            if current_num is not None or current_parts:
                entries.append((current_num, " ".join(current_parts)))
            current_num = int(match.group("num"))
            current_parts = [match.group("name").strip().rstrip(".")]
            numbered_seen += 1
            continue

        if current_num is not None and not _looks_like_numbered_list_boundary(line):
            if _is_numbered_list_section_header(line):
                entries.append((current_num, " ".join(current_parts)))
                current_num = None
                current_parts = []
                continue
            if numbered_seen >= 3 and _looks_like_unnumbered_list_item(line):
                entries.append((current_num, " ".join(current_parts)))
                entries.append((None, line.rstrip(".")))
                current_num = None
                current_parts = []
                continue
            current_parts.append(line.rstrip("."))
            continue

        if current_parts:
            entries.append((current_num, " ".join(current_parts)))
            current_num = None
            current_parts = []

        if numbered_seen >= 3 and _looks_like_unnumbered_list_item(line):
            entries.append((None, line.rstrip(".")))

    if current_parts:
        entries.append((current_num, " ".join(current_parts)))

    return entries


def _looks_like_numbered_list_boundary(line: str) -> bool:
    lower = line.lower()
    if _NUMBERED_LIST_ITEM_RE.match(line):
        return True
    if any(marker in lower for marker in _NUMBERED_LIST_SKIP_MARKERS):
        return True
    if lower.startswith("приложение"):
        return True
    return False


def _looks_like_unnumbered_list_item(line: str) -> bool:
    if _NUMBERED_LIST_ITEM_RE.match(line):
        return False
    if _is_numbered_list_section_header(line):
        return False
    if len(line) < 12:
        return False
    lower = line.lower()
    if any(marker in lower for marker in _NUMBERED_LIST_SKIP_MARKERS):
        return False
    if lower.startswith("приложение"):
        return False
    if not re.search(r"[а-яa-zё]{4,}", lower):
        return False
    return _is_valid_numbered_list_name(line)


def _parse_numbered_list_text(text: str) -> list[TZItem]:
    if not text.strip():
        return []

    text = _repair_ocr_cyrillic_text(text)
    if not _looks_like_numbered_equipment_list(text):
        return []

    items: list[TZItem] = []
    seen_numbers: set[int] = set()
    auto_number = 0

    for raw_number, raw_name in _collect_numbered_list_entries(text):
        name = _normalize_tz_product_name(raw_name)
        if not _is_valid_numbered_list_name(name):
            continue

        if raw_number is not None and raw_number > 0 and raw_number not in seen_numbers:
            number = raw_number
            seen_numbers.add(number)
            auto_number = max(auto_number, number)
        else:
            auto_number += 1
            number = auto_number

        items.append(
            TZItem(
                number=number,
                name=name,
                unit="шт.",
                quantity=1.0,
            )
        )

    if len(items) < 3:
        return []

    return items


def _looks_like_procurement_cover_letter(text: str) -> bool:
    normalized = _repair_ocr_cyrillic_text(text).lower()
    markers = (
        "запрос",
        "3апрос",
        "zampoc",
        "коммерч",
        "kouмep",
        "koммep",
        "коммерческ",
        "приложен",
        "прилоx",
        "npu nox",
        "npunox",
        "ценов",
        "ценовой",
        "mpeдlo",
        "предложen",
        "предложен",
        "mpeocma",
        "mpeдocma",
    )
    org_markers = (
        "учтендер",
        "yutender",
        "yuten",
        "uten",
        "уten",
        "казенное",
        "образovan",
        "образован",
        "закупк",
        "аукцион",
    )
    marker_hits = sum(1 for marker in markers if marker in normalized)
    org_hits = sum(1 for marker in org_markers if marker in normalized)
    if marker_hits >= 2 or (marker_hits >= 1 and org_hits >= 1) or org_hits >= 2:
        return not _looks_like_numbered_equipment_list(text)
    return False


def _ocr_text_looks_like_official_letter(text: str) -> bool:
    if not text.strip():
        return False
    upper_lines = 0
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if len(line) < 18:
            continue
        letters = [char for char in line if char.isalpha()]
        if len(letters) < 12:
            continue
        upper_count = sum(1 for char in letters if char.isupper())
        if upper_count / len(letters) >= 0.82:
            upper_lines += 1
    return upper_lines >= 3


def _parse_price_request_tables(tables: Iterable[list[list[str]]]) -> list[TZItem]:
    items: list[TZItem] = []
    for table in tables:
        if len(table) < 2:
            continue
        header_cells = [_cell_str(cell) for cell in table[0]]
        joined = " ".join(cell.strip().lower() for cell in header_cells if cell.strip())
        if not _looks_like_price_request_header(joined):
            continue
        col = _build_price_request_column_map(header_cells)
        if not col:
            continue

        auto_number = 0
        for row in table[1:]:
            cells = [_cell_str(cell) for cell in row]
            if len(cells) <= col.name:
                continue
            number_raw = cells[col.number].rstrip(".")
            if number_raw.isdigit():
                number = int(number_raw)
            else:
                auto_number += 1
                number = auto_number
            row_text = " | ".join(cells)
            vendor_match = _VENDOR_CODE_RE.search(row_text)
            vendor_code = vendor_match.group(1).upper() if vendor_match else ""
            raw_name = _normalize_tz_product_name(cells[col.name])
            name = raw_name
            if vendor_code and vendor_code not in raw_name.upper():
                name = f"{vendor_code} {raw_name}".strip()
            if not _is_valid_tz_product_name(name):
                continue

            quantity = 1.0
            if col.qty is not None and len(cells) > col.qty:
                parsed_qty = _parse_quantity(cells[col.qty])
                if parsed_qty is not None and parsed_qty > 0:
                    quantity = parsed_qty

            specs = f"Код производителя: {vendor_code}" if vendor_code else ""
            items.append(
                TZItem(
                    number=number,
                    name=name,
                    unit="шт.",
                    quantity=quantity,
                    specifications=specs,
                )
            )
    return items


_EIS_KNOWN_CHARACTERISTICS: tuple[str, ...] = (
    "Конструкционные особенности",
    "Строение оптической схемы",
    "Максимальное увеличение",
    "Разрешение камеры",
    "Расположение осветителя",
    "Разъем входа/выхода",
    "Способ наблюдения",
    "Тип осветителя",
    "Тип матрицы",
    "Питание",
)


def _eis_characteristics_pattern() -> str:
    return "|".join(
        re.escape(header)
        for header in sorted(_EIS_KNOWN_CHARACTERISTICS, key=len, reverse=True)
    )


def _is_eis_tz_document(text: str) -> bool:
    lower = text.lower()
    if "наимен" not in lower and "техническое задание" not in lower:
        return False
    if "характер" not in lower and "осветител" not in lower and "микроскоп" not in lower:
        return False
    if "\x07" in text and re.search(r"\d+\x07", text):
        return True
    if re.search(r"\d+\s*\|", text):
        return True
    if "\t" in text and re.search(r"\d+\t", text):
        return True
    if re.search(r"\b1\s+[А-Яа-яA-Za-z«»]", text) and re.search(
        r"\bШТ\s*\d+", text, re.IGNORECASE
    ):
        return True
    if re.search(r"\b1[А-Яа-яA-Za-z«»]", text) and re.search(
        r"ШТ\s*\d+", text, re.IGNORECASE
    ):
        return True
    return False


def _clean_eis_char_value(raw: str) -> str:
    cleaned = re.sub(
        r"\b(ШТ|ШТ\.|компл|упак)\s*\d+\s*(?:Российская Федерация)?",
        "",
        raw,
        flags=re.IGNORECASE,
    )
    return re.sub(r"\s+", " ", cleaned).strip()


def _is_readable_spec_fragment(text: str) -> bool:
    cleaned = text.strip()
    if not cleaned:
        return False
    if "\x00" in cleaned:
        return False
    if re.search(r"[\u0900-\u0dff\u0e00-\u0eff\uf000-\uffff]", cleaned):
        return False
    if re.fullmatch(r"[\d.,]+", cleaned):
        return True
    letters = len(re.findall(r"[А-Яа-яA-Za-zЁё]", cleaned))
    return letters >= max(2, len(cleaned) // 4)


def _is_eis_characteristic_header(text: str) -> bool:
    cleaned = text.strip()
    if not cleaned:
        return False
    lower = cleaned.lower()
    for header in _EIS_KNOWN_CHARACTERISTICS:
        if lower == header.lower():
            return True
    return _looks_like_characteristic_name(cleaned)


def _sanitize_eis_delimited_text(text: str) -> str:
    if "\x07" not in text:
        return text
    text = text.replace("\x00", "")
    text = re.sub(
        r"[^\w\s«»\-.,/():;№%\"'А-Яа-яЁё\x07]+",
        "\x07",
        text,
    )
    return re.sub(r"\x07{4,}", "\x07\x07", text)


def _extract_eis_specifications(segment: str) -> list[str]:
    pattern = re.compile(
        rf"({_eis_characteristics_pattern()})",
        re.IGNORECASE,
    )
    matches = list(pattern.finditer(segment))
    spec_lines: list[str] = []
    for index, match in enumerate(matches):
        header = match.group(1).strip()
        value_start = match.end()
        value_end = (
            matches[index + 1].start() if index + 1 < len(matches) else len(segment)
        )
        value = _clean_eis_char_value(segment[value_start:value_end])
        if value and _is_readable_spec_fragment(value):
            spec_lines.append(f"{header}: {value}")
    return spec_lines


def _eis_data_section(text: str) -> str:
    normalized = _normalize_eis_doc_text(text)
    lower = normalized.lower()
    data_start = 0
    for marker in (
        "единица измерения характеристики",
        "наименование характеристики",
        "кол-во товара",
        "наименование страны происхождения товара",
    ):
        idx = lower.find(marker)
        if idx >= 0:
            data_start = max(data_start, idx + len(marker))
    return normalized[data_start:].strip() or normalized


def _parse_zakupki_concatenated_text(text: str) -> list[TZItem]:
    if not _is_eis_tz_document(text):
        return []

    scan = _eis_data_section(text)
    if not scan:
        return []

    headers_alt = _eis_characteristics_pattern()
    items: list[TZItem] = []
    item_pattern = re.compile(
        rf"\b(\d+)\s+([А-Яа-яA-Za-z«»\"][^|]+?)(?=\s+(?:{headers_alt}))",
        re.IGNORECASE,
    )
    for match in item_pattern.finditer(scan):
        number = int(match.group(1))
        if number <= 0 or number > 200:
            continue

        name = _normalize_tz_product_name(match.group(2))
        if not _is_valid_tz_product_name(name):
            continue

        tail = scan[match.start() :]
        segment_end = len(tail)
        itogo = re.search(r"\bИТОГО\b", tail, re.IGNORECASE)
        if itogo:
            segment_end = itogo.start()

        for later in item_pattern.finditer(tail[match.end() - match.start() :]):
            if later.start() <= 0:
                continue
            later_number = int(later.group(1))
            later_name = _normalize_tz_product_name(later.group(2))
            if later_number <= 0 or later_number > 200:
                continue
            if not _is_valid_tz_product_name(later_name):
                continue
            later_tail = tail[match.end() - match.start() + later.start() :]
            if re.search(r"\bШТ\s*\d+", later_tail[:400], re.IGNORECASE):
                segment_end = min(segment_end, match.end() - match.start() + later.start())
                break

        segment = tail[:segment_end]

        if not re.search(r"\bШТ\s*\d+", segment, re.IGNORECASE):
            continue

        unit = "шт."
        quantity = 1.0
        country = ""
        meta = re.search(
            r"\b(ШТ|ШТ\.|компл|упак)\s*(\d+(?:[.,]\d+)?)\s*(Российская Федерация)?",
            segment,
            re.IGNORECASE,
        )
        if meta:
            unit = meta.group(1)
            quantity = float(meta.group(2).replace(",", "."))
            if meta.group(3):
                country = meta.group(3).strip()

        spec_lines = _extract_eis_specifications(segment)
        if not spec_lines:
            continue

        items.append(
            TZItem(
                number=number,
                name=name,
                unit=unit,
                quantity=quantity,
                specifications="; ".join(spec_lines),
                country_of_origin=country,
            )
        )

    return items


def _build_column_map(header_cells: list[str]) -> _ColumnMap | None:
    lower = [c.strip().lower() for c in header_cells]
    if not any("наимен" in cell for cell in lower):
        return None

    country_idx: int | None = None
    for i, cell in enumerate(lower):
        if "стран" in cell or "происхожд" in cell:
            country_idx = i
            break

    name_idx: int | None = None
    for i, cell in enumerate(lower):
        if "наимен" not in cell:
            continue
        if "стран" in cell or "происхожд" in cell:
            continue
        name_idx = i
        break
    if name_idx is None:
        return None

    number_idx = 0
    has_number = False
    for i, cell in enumerate(lower):
        if cell in {"№", "п/п", "пп", "n", "no"} or cell.startswith("№"):
            number_idx = i
            has_number = True
            break
        if "номер" in cell and "товар" not in cell and "стран" not in cell:
            number_idx = i
            has_number = True
            break

    unit_idx = name_idx + 1
    for i, cell in enumerate(lower):
        if ("ед" in cell and ("изм" in cell or cell.startswith("ед"))) or "единиц" in cell:
            unit_idx = i
            break

    qty_idx = name_idx + 2
    for i, cell in enumerate(lower):
        if "кол" in cell or "колич" in cell:
            qty_idx = i
            break

    specs_idx: int | None = None
    for i, cell in enumerate(lower):
        if any(key in cell for key in ("характер", "описан", "функцион", "техн")):
            specs_idx = i
            break

    sale_price_idx: int | None = None
    sale_price_priority = -1
    reserved_cols = {number_idx, name_idx, unit_idx, qty_idx}
    if specs_idx is not None:
        reserved_cols.add(specs_idx)
    if country_idx is not None:
        reserved_cols.add(country_idx)
    for i, cell in enumerate(lower):
        if i in reserved_cols:
            continue
        if any(
            marker in cell
            for marker in ("себестоим", "закуп", "поставщ", "ндс", "итого")
        ):
            continue
        if "сумма" in cell and "цена" not in cell:
            continue
        priority = -1
        if any(marker in cell for marker in ("продаж", "отпуск", "реализац")):
            priority = 3
        elif "цена" in cell and ("ед" in cell or "единиц" in cell):
            priority = 2
        elif "цена" in cell:
            priority = 1
        elif "стоимость" in cell and ("ед" in cell or "единиц" in cell):
            priority = 1
        if priority > sale_price_priority:
            sale_price_priority = priority
            sale_price_idx = i

    return _ColumnMap(
        number_idx,
        name_idx,
        unit_idx,
        qty_idx,
        specs_idx,
        country_idx,
        sale_price_idx,
        has_number,
    )


def _parse_row(
    cells: list[str],
    col: _ColumnMap,
    *,
    fallback_number: int | None = None,
) -> TZItem | None:
    indices = [col.name, col.unit, col.qty]
    if col.specs is not None:
        indices.append(col.specs)
    if col.country is not None:
        indices.append(col.country)
    max_idx = max(indices)
    if len(cells) <= max_idx:
        return None

    number: int | None = None
    if col.has_number:
        number_raw = cells[col.number].rstrip(".")
        if number_raw.isdigit():
            number = int(number_raw)
        elif fallback_number is not None:
            number = fallback_number
        else:
            return None
    elif fallback_number is not None:
        number = fallback_number
    else:
        return None

    name = _normalize_tz_product_name(cells[col.name])
    if not _is_valid_tz_product_name(name):
        return None

    unit = cells[col.unit] or "шт."
    qty_raw = cells[col.qty].replace(",", ".")
    specs = cells[col.specs] if col.specs is not None and len(cells) > col.specs else ""
    country = (
        cells[col.country]
        if col.country is not None and len(cells) > col.country
        else ""
    )

    try:
        quantity = float(qty_raw)
    except ValueError:
        quantity = 1.0

    sale_price = None
    if col.sale_price is not None and len(cells) > col.sale_price:
        sale_price = _parse_money(cells[col.sale_price])
        if (
            sale_price is not None
            and number is not None
            and sale_price == float(number)
            and sale_price < 1000
        ):
            sale_price = None

    return TZItem(
        number=number,
        name=name,
        unit=unit,
        quantity=quantity,
        specifications=specs,
        country_of_origin=country,
        target_sale_price=sale_price,
    )


_SPEC_UNIT_TOKENS = frozenset(
    {
        "шт",
        "шт.",
        "штука",
        "штук",
        "компл",
        "компл.",
        "комплект",
        "упак",
        "упак.",
        "мм",
        "см",
        "м",
        "кг",
        "г",
        "л",
        "мл",
        "вт",
        "квт",
        "об/мин",
        "об",
        "мин",
        "час",
        "°с",
        "°c",
        "мпа",
        "гц",
        "дб",
        "%",
    }
)


def _looks_like_spec_unit(value: str) -> bool:
    normalized = value.strip().lower().rstrip(".")
    if not normalized:
        return False
    if normalized in _SPEC_UNIT_TOKENS:
        return True
    if normalized in _PRODUCT_UNITS:
        return True
    if re.fullmatch(r"[\d°%./\-a-zа-яё]{1,12}", normalized):
        return len(normalized) <= 8
    return False


def _dominant_product_name_column(rows: list[list[str]]) -> tuple[int, str] | None:
    if len(rows) < 2:
        return None

    col_count = max(len(row) for row in rows)
    best: tuple[int, str] | None = None
    best_score = 0.0

    for col_idx in range(col_count):
        values: list[str] = []
        for row in rows:
            if col_idx >= len(row):
                continue
            normalized = _normalize_tz_product_name(_cell_str(row[col_idx]))
            if normalized:
                values.append(normalized)
        if len(values) < 2:
            continue

        counts: dict[str, int] = {}
        for value in values:
            counts[value] = counts.get(value, 0) + 1
        name, repeats = max(counts.items(), key=lambda item: item[1])
        ratio = repeats / len(rows)
        if ratio < 0.55:
            continue
        if not _is_valid_tz_product_name(name):
            continue
        score = ratio * len(name)
        if score > best_score:
            best_score = score
            best = (col_idx, name)

    return best


def _detect_vertical_spec_columns(
    rows: list[list[str]],
) -> tuple[int, int, int, int | None] | None:
    dominant = _dominant_product_name_column(rows)
    if not dominant:
        return None

    name_col, _ = dominant
    col_count = max(len(row) for row in rows)
    if col_count < 3:
        return None

    best_param_col: int | None = None
    best_param_score = 0
    for col_idx in range(col_count):
        if col_idx == name_col:
            continue
        score = 0
        for row in rows:
            if col_idx >= len(row):
                continue
            text = _cell_str(row[col_idx])
            if not text or len(text) < 3:
                continue
            if text.replace(" ", "").isdigit():
                continue
            if _looks_like_characteristic_name(text) or len(text) <= 120:
                score += 1
        if score > best_param_score:
            best_param_score = score
            best_param_col = col_idx

    if best_param_col is None or best_param_score < 2:
        return None

    value_col: int | None = None
    value_score = 0
    for col_idx in range(col_count):
        if col_idx in {name_col, best_param_col}:
            continue
        score = sum(
            1
            for row in rows
            if col_idx < len(row) and _cell_str(row[col_idx]).strip()
        )
        if score > value_score:
            value_score = score
            value_col = col_idx

    if value_col is None or value_score < 2:
        return None

    unit_col: int | None = None
    for col_idx in range(col_count):
        if col_idx in {name_col, best_param_col, value_col}:
            continue
        unit_hits = sum(
            1
            for row in rows
            if col_idx < len(row) and _looks_like_spec_unit(_cell_str(row[col_idx]))
        )
        if unit_hits >= max(2, len(rows) // 4):
            unit_col = col_idx
            break

    return name_col, best_param_col, value_col, unit_col


def _is_vertical_spec_table(rows: list[list[str]]) -> bool:
    if len(rows) < 3:
        return False
    if _is_tz_table_header([_cell_str(cell) for cell in rows[0]]):
        return False
    return _detect_vertical_spec_columns(rows) is not None


def _format_vertical_spec_line(param: str, value: str, unit: str) -> str:
    param = re.sub(r"\s+", " ", param.replace("\n", " ")).strip(" :")
    value = re.sub(r"\s+", " ", value.replace("\n", " ")).strip()
    unit = re.sub(r"\s+", " ", unit.replace("\n", " ")).strip()
    if not param:
        return ""
    if not value:
        return param
    if unit and unit.lower() not in value.lower():
        return f"{param}: {value} {unit}".strip()
    return f"{param}: {value}"


def _parse_vertical_spec_table(
    rows: list[list[str]],
    *,
    number: int,
) -> TZItem | None:
    columns = _detect_vertical_spec_columns(rows)
    if not columns:
        return None

    name_col, param_col, value_col, unit_col = columns
    dominant = _dominant_product_name_column(rows)
    if not dominant:
        return None

    _, name = dominant
    spec_lines: list[str] = []
    for row in rows:
        cells = [_cell_str(cell) for cell in row]
        if param_col >= len(cells):
            continue
        param = cells[param_col]
        value = cells[value_col] if value_col < len(cells) else ""
        unit = cells[unit_col] if unit_col is not None and unit_col < len(cells) else ""
        line = _format_vertical_spec_line(param, value, unit)
        if line:
            spec_lines.append(line)

    if not spec_lines:
        return None

    return TZItem(
        number=number,
        name=name,
        unit="шт.",
        quantity=1.0,
        specifications="\n".join(spec_lines),
    )


def _parse_vertical_spec_tables(tables: Iterable[list[list[str]]]) -> list[TZItem]:
    items: list[TZItem] = []
    for table in tables:
        if not _is_vertical_spec_table(table):
            continue
        item = _parse_vertical_spec_table(table, number=len(items) + 1)
        if item:
            items.append(item)
    return items


def _parse_tz_tables(tables: Iterable[list[list[str]]]) -> list[TZItem]:
    items = _parse_price_request_tables(tables)
    if items:
        return items

    items: list[TZItem] = []

    for table in tables:
        if len(table) < 2:
            continue

        header_cells = [_cell_str(cell) for cell in table[0]]
        if not _is_tz_table_header(header_cells):
            continue
        col = _build_column_map(header_cells)
        if not col:
            continue

        auto_number = 0
        for row in table[1:]:
            cells = [_cell_str(cell) for cell in row]
            if not any(cells):
                continue
            auto_number += 1
            item = _parse_row(cells, col, fallback_number=auto_number)
            if item:
                items.append(item)

    return items


def _detect_ole_suffix(content: bytes) -> str:
    if b"Workbook" in content and b"MSWord" not in content and b"WordDocument" not in content:
        return ".xls"
    if b"MSWord" in content or b"WordDocument" in content:
        return ".doc"
    if b"Word" in content and b"Excel" not in content[:8192]:
        return ".doc"
    return ".xls"


def _extract_doc_text_from_ole_utf16(path: Path) -> str:
    """Извлекает текст из старого .doc напрямую из UTF-16LE-потока (без textutil/antiword)."""
    data = path.read_bytes()
    if not data.startswith(b"\xd0\xcf\x11\xe0"):
        return ""

    start = -1
    for label in ("Техническое задание", "Наименование", "наименование"):
        idx = data.find(label.encode("utf-16le"))
        if idx >= 0:
            start = idx if start < 0 else min(start, idx)
    if start < 0:
        return ""
    if start % 2 == 1:
        start += 1

    end = -1
    for end_label in ("ИТОГО",):
        idx = data.find(end_label.encode("utf-16le"), start)
        if idx >= 0:
            end = idx + len(end_label.encode("utf-16le")) + 400
            break
    if end < 0:
        end = min(len(data), start + 80000)

    blob = data[start:end]
    if len(blob) % 2 == 1:
        blob = blob[:-1]
    return blob.decode("utf-16le", errors="ignore").strip()


def _score_doc_extraction_text(text: str) -> int:
    if not text.strip():
        return 0
    lower = text.lower()
    score = len(re.findall(r"[а-яё]", lower))
    if "\x07" in text:
        score += 50
    if "наимен" in lower:
        score += 20
    if re.search(r"шт\s*\d+", lower):
        score += 20
    if "характер" in lower:
        score += 10
    return score


def _normalize_eis_doc_text(text: str) -> str:
    normalized = text.replace("\f", " ")
    normalized = re.sub(r"[\r\n\t]+", " ", normalized)
    normalized = re.sub(r"\s*\|\s*", " ", normalized)
    return re.sub(r"\s+", " ", normalized).strip()


def _parse_zakupki_loose_text(text: str) -> list[TZItem]:
    normalized = _normalize_eis_doc_text(text)
    if not normalized:
        return []

    spec_lines = _extract_eis_specifications(normalized)
    if len(spec_lines) < 1:
        return []

    char_pattern = re.compile(rf"({_eis_characteristics_pattern()})", re.IGNORECASE)
    first_char = char_pattern.search(normalized)
    if not first_char:
        return []

    prefix = normalized[: first_char.start()]
    name: str | None = None
    number = 1
    for match in reversed(
        list(
            re.finditer(
                r"(?:^|[\s\x07])(\d+)\s*[\.\):\-–—]?\s*"
                r"([А-Яа-яA-Za-z«»\"][\w\s«»\-]{4,})",
                prefix,
                re.IGNORECASE,
            )
        )
    ):
        candidate_number = int(match.group(1))
        if candidate_number <= 0 or candidate_number > 200:
            continue
        candidate_name = _normalize_tz_product_name(match.group(2))
        if not _is_valid_tz_product_name(candidate_name):
            continue
        number = candidate_number
        name = candidate_name
        break

    if not name:
        return []

    unit = "шт."
    quantity = 1.0
    country = ""
    meta = re.search(
        r"\b(ШТ|ШТ\.|компл|упак)\s*(\d+(?:[.,]\d+)?)\s*(Российская Федерация)?",
        normalized,
        re.IGNORECASE,
    )
    if meta:
        unit = meta.group(1)
        quantity = float(meta.group(2).replace(",", "."))
        if meta.group(3):
            country = meta.group(3).strip()

    return [
        TZItem(
            number=number,
            name=name,
            unit=unit,
            quantity=quantity,
            specifications="; ".join(spec_lines),
            country_of_origin=country,
        )
    ]


def _extract_doc_text(path: Path) -> str:
    errors: list[str] = []
    candidates: list[str] = []

    ole_text = _extract_doc_text_from_ole_utf16(path)
    if ole_text:
        candidates.append(ole_text)

    if shutil.which("textutil"):
        try:
            candidates.append(
                subprocess.check_output(
                    ["textutil", "-convert", "txt", "-stdout", str(path)],
                    stderr=subprocess.STDOUT,
                ).decode("utf-8", errors="replace")
            )
        except subprocess.CalledProcessError as exc:
            errors.append(f"textutil: {exc}")

    if shutil.which("antiword"):
        for args in (
            ["-m", "UTF-8.txt"],
            ["-m", "UTF-8.txt", "-w", "0"],
            [],
        ):
            try:
                raw = subprocess.check_output(
                    ["antiword", *args, str(path)],
                    stderr=subprocess.STDOUT,
                )
                candidates.append(raw.decode("utf-8", errors="replace"))
                for encoding in ("cp1251", "cp866"):
                    candidates.append(raw.decode(encoding, errors="replace"))
            except subprocess.CalledProcessError as exc:
                errors.append(f"antiword{' '.join(args)}: {exc}")

    if shutil.which("soffice"):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                result = subprocess.run(
                    [
                        "soffice",
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        tmp,
                        str(path),
                    ],
                    check=False,
                    capture_output=True,
                    text=True,
                )
                if result.returncode != 0:
                    errors.append(f"soffice: {result.stderr.strip() or result.stdout.strip()}")
                else:
                    txt_path = Path(tmp) / f"{path.stem}.txt"
                    if txt_path.exists():
                        candidates.append(
                            txt_path.read_text(encoding="utf-8", errors="replace")
                        )
        except OSError as exc:
            errors.append(f"soffice: {exc}")

    candidates = [text for text in candidates if text and text.strip()]
    if candidates:
        return max(candidates, key=_score_doc_extraction_text)

    details = f" ({'; '.join(errors)})" if errors else ""
    raise ValueError(
        "Не удалось прочитать .doc. Сохраните файл как .docx "
        f"или установите LibreOffice/antiword.{details}"
    )


def _parse_tz_doc(path: Path) -> list[TZItem]:
    text = _extract_doc_text(path)
    items = _parse_zakupki_doc_stream(text)
    if items:
        return items

    items = _parse_zakupki_loose_text(text)
    if items:
        return items

    items = _parse_tz_tables(_tables_from_text(text))
    if items:
        return items

    raise ValueError(
        "Не удалось извлечь позиции из .doc. "
        "Убедитесь, что в файле есть таблица с колонкой «Наименование товара»."
    )


def _is_product_unit(value: str) -> bool:
    normalized = value.strip().lower().rstrip(".")
    if normalized in _PRODUCT_UNITS:
        return True
    upper = value.strip().upper()
    return upper in {"ШТ", "КОМПЛ", "УПАК"}


def _parse_quantity(value: str) -> float | None:
    raw = value.strip().replace(" ", "").replace(",", ".")
    if not raw:
        return None
    try:
        return float(raw)
    except ValueError:
        return None


def _parse_money(value: str) -> float | None:
    raw = str(value or "").strip()
    if not raw or raw in {"—", "-", "–"}:
        return None
    normalized = (
        raw.lower()
        .replace("руб.", "")
        .replace("руб", "")
        .replace("₽", "")
        .replace(" ", "")
        .replace("\u00a0", "")
        .replace(",", ".")
    )
    normalized = re.sub(r"[^0-9.\-]", "", normalized)
    if not normalized or normalized in {".", "-", "-."}:
        return None
    try:
        amount = float(normalized)
    except ValueError:
        return None
    return amount if amount > 0 else None


def _looks_like_characteristic_name(value: str) -> bool:
    lower = value.lower()
    keywords = (
        "располож",
        "тип",
        "способ",
        "размер",
        "максим",
        "миним",
        "конструк",
        "разреш",
        "разъем",
        "питание",
        "строение",
        "матриц",
        "характер",
        "назнач",
        "мощност",
        "диаметр",
        "длина",
        "ширин",
        "высот",
        "вес",
        "объем",
        "емкост",
        "частот",
        "напряж",
        "ротор",
        "таймер",
        "скорост",
        "комплект",
        "управл",
        "режим",
        "температ",
        "вместим",
        "дискрет",
        "функц",
        "поддерж",
        "адаптер",
        "центрифуг",
    )
    return any(keyword in lower for keyword in keywords)


def _parse_zakupki_item_parts(parts: list[str]) -> TZItem | None:
    if len(parts) < 2:
        return None

    number_raw = parts[0].rstrip(".")
    if not number_raw.isdigit():
        return None

    name = _normalize_tz_product_name(parts[1])
    if not _is_valid_tz_product_name(name):
        return None

    unit = "шт."
    quantity = 1.0
    country = ""
    spec_lines: list[str] = []
    product_meta_found = False

    idx = 2
    while idx < len(parts):
        part = parts[idx].strip()
        if not part:
            idx += 1
            continue
        if part.upper().startswith("ИТОГО"):
            break

        if not product_meta_found and _is_product_unit(part):
            qty = _parse_quantity(parts[idx + 1]) if idx + 1 < len(parts) else None
            if qty is not None:
                unit = part.strip() or "шт."
                quantity = qty
                idx += 2
                if idx < len(parts):
                    candidate = parts[idx].strip()
                    if (
                        candidate
                        and not _is_product_unit(candidate)
                        and _parse_quantity(candidate) is None
                        and (
                            not _looks_like_characteristic_name(candidate)
                            or len(candidate) > 15
                        )
                    ):
                        country = candidate
                        idx += 1
                product_meta_found = True
                continue

        if (
            idx + 1 < len(parts)
            and _is_eis_characteristic_header(part)
            and _is_readable_spec_fragment(part)
        ):
            value = parts[idx + 1].strip()
            if (
                value
                and not value.upper().startswith("ИТОГО")
                and _is_readable_spec_fragment(value)
            ):
                spec_value = value
                step = 2
                if idx + 2 < len(parts):
                    char_unit = parts[idx + 2].strip()
                    if (
                        char_unit
                        and not _is_product_unit(char_unit)
                        and _parse_quantity(char_unit) is None
                        and len(char_unit) <= 20
                        and not _looks_like_characteristic_name(char_unit)
                        and _is_readable_spec_fragment(char_unit)
                    ):
                        spec_value = f"{value} {char_unit}"
                        step = 3
                spec_lines.append(f"{part}: {spec_value}")
                idx += step
                continue

        idx += 1

    if not spec_lines and not product_meta_found:
        return None

    return TZItem(
        number=int(number_raw),
        name=name,
        unit=unit,
        quantity=quantity,
        specifications="; ".join(spec_lines),
        country_of_origin=country,
    )


def _iter_zakupki_product_matches(
    text: str,
    delimiter: str,
) -> list[re.Match[str]]:
    if delimiter == "\x07":
        if "\x07" not in text:
            return []
        item_pattern = re.compile(
            r"(?:^|[\n\r]|\x07)(\d+)\x07"
            r"([А-Яа-яA-Za-z«»\"][^\x07\n\r]{2,})",
            re.IGNORECASE,
        )
    elif delimiter == "|":
        if not re.search(r"\d+\s*\|", text) and not re.search(r"\|\s*\d+\s*\|", text):
            return []
        item_pattern = re.compile(
            r"(?:^|[\n\r]|\|)\s*(\d+)\s*\|"
            r"([А-Яа-яA-Za-z«»\"][^|\n\r]{2,})",
            re.IGNORECASE,
        )
    else:
        if delimiter not in text:
            return []
        item_pattern = re.compile(
            rf"(?:^|[\n\r]|\|)\s*(\d+)\s*{re.escape(delimiter)}"
            rf"([А-Яа-яA-Za-z«»\"][^{re.escape(delimiter)}\n\r]{{2,}})",
            re.IGNORECASE,
        )

    matches: list[re.Match[str]] = []
    for match in item_pattern.finditer(text):
        number = int(match.group(1))
        if number <= 0 or number > 200:
            continue
        name = _normalize_tz_product_name(match.group(2))
        if not _is_valid_tz_product_name(name):
            continue
        matches.append(match)
    return matches


def _parse_zakupki_delimited_stream(text: str, delimiter: str) -> list[TZItem]:
    if delimiter == "\x07":
        split_delimiter = "\x07"
    elif delimiter == "|":
        split_delimiter = "|"
    else:
        split_delimiter = delimiter

    row_matches = _iter_zakupki_product_matches(text, delimiter)
    if not row_matches:
        return []

    items: list[TZItem] = []
    for index, match in enumerate(row_matches):
        tail = text[match.start(1) :]
        segment_end = len(tail)
        itogo = re.search(r"\bИТОГО\b", tail, re.IGNORECASE)
        if itogo:
            segment_end = itogo.start()
        if index + 1 < len(row_matches):
            next_start = row_matches[index + 1].start(1) - match.start(1)
            if next_start > 0:
                segment_end = min(segment_end, next_start)

        parts = [part.strip() for part in tail[:segment_end].split(split_delimiter)]
        item = _parse_zakupki_item_parts(parts)
        if item:
            items.append(item)

    return items


def _zakupki_items_complete(items: list[TZItem], source_text: str) -> bool:
    if not items:
        return False
    for item in items:
        if not item.specifications or len(item.specifications) < 40:
            return False
        if item.quantity <= 1 and re.search(r"\bШТ\s+[2-9]\d", source_text, re.IGNORECASE):
            return False
    return True


def _parse_zakupki_doc_stream(text: str) -> list[TZItem]:
    if not _is_eis_tz_document(text):
        return []

    if "\x07" in text:
        text = _sanitize_eis_delimited_text(text)

    best_delimited: list[TZItem] = []
    for delimiter in ("\x07", "|", "\t"):
        items = _parse_zakupki_delimited_stream(text, delimiter)
        if items:
            if _zakupki_items_complete(items, text):
                return items
            best_delimited = items

    items = _parse_zakupki_concatenated_text(text)
    if items:
        return items

    items = _parse_zakupki_loose_text(text)
    if items:
        return items

    return best_delimited


def _parse_zakupki_doc_cell_stream(text: str) -> list[TZItem]:
    return _parse_zakupki_doc_stream(text)


def _parse_tz_docx(path: Path) -> list[TZItem]:
    doc = Document(str(path))
    tables: list[list[list[str]]] = []

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(rows)

    items = _parse_tz_tables(tables)
    if items:
        return items

    items = _parse_vertical_spec_tables(tables)
    if items:
        return items

    text = _extract_docx_text(path)
    items = _parse_numbered_list_text(text)
    if items:
        return items

    items = _parse_zakupki_loose_text(text)
    if items:
        return items

    return []


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    chunks: list[str] = []
    chunks.extend(p.text.strip() for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                chunks.append(" | ".join(values))
    return "\n".join(chunks)


def _parse_tz_excel(path: Path) -> list[TZItem]:
    suffix = path.suffix.lower()
    sheets: list[list[list[str]]] = []

    if suffix == ".xlsx":
        wb = load_workbook(path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            sheet_rows: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                sheet_rows.append([_cell_str(cell) for cell in row])
            if sheet_rows:
                sheets.append(sheet_rows)
        wb.close()
    else:
        wb = xlrd.open_workbook(str(path))
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            sheet_rows = [
                [_cell_str(cell) for cell in ws.row_values(row_idx)]
                for row_idx in range(ws.nrows)
            ]
            if sheet_rows:
                sheets.append(sheet_rows)

    tables = [_table_from_sheet_rows(sheet) for sheet in sheets]
    tables = [table for table in tables if table]
    return _parse_tz_tables(tables)


def _extract_excel_text(path: Path) -> str:
    suffix = path.suffix.lower()
    lines: list[str] = []
    if suffix == ".xlsx":
        wb = load_workbook(path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            lines.append(f"# Лист: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = [_cell_str(cell) for cell in row if _cell_str(cell)]
                if values:
                    lines.append(" | ".join(values))
        wb.close()
    else:
        wb = xlrd.open_workbook(str(path))
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            lines.append(f"# Лист: {sheet_name}")
            for row_idx in range(ws.nrows):
                values = [_cell_str(cell) for cell in ws.row_values(row_idx)]
                values = [value for value in values if value]
                if values:
                    lines.append(" | ".join(values))
    return "\n".join(lines)


def _table_from_sheet_rows(sheet_rows: list[list[str]]) -> list[list[str]] | None:
    for idx, row in enumerate(sheet_rows):
        if _build_column_map(row):
            return sheet_rows[idx:]
    return None


def _pdf_page_is_image_only(path: Path) -> bool:
    try:
        import fitz
    except ImportError:
        return False

    try:
        with fitz.open(path) as doc:
            saw_image = False
            for page in doc:
                if page.get_text("text").strip():
                    return False
                blocks = page.get_text("dict").get("blocks") or []
                for block in blocks:
                    if block.get("type") == 0:
                        return False
                    if block.get("type") == 1:
                        saw_image = True
            return saw_image
    except Exception:
        logger.debug("Failed to inspect PDF structure %s", path, exc_info=True)
        return False


_PDF_TEXT_CACHE: dict[str, str] = {}


def _extract_pdf_full_text(path: Path) -> str:
    cache_key = str(path.resolve())
    cached = _PDF_TEXT_CACHE.get(cache_key)
    if cached is not None:
        return cached

    text = _extract_pdf_text(path)
    if not text.strip() and TZ_PDF_OCR_ENABLED and _pdf_page_is_image_only(path):
        text = _ocr_pdf(path)
    elif not text.strip() and TZ_PDF_OCR_ENABLED:
        text = _ocr_pdf(path)

    text = _repair_ocr_cyrillic_text(text)
    _PDF_TEXT_CACHE[cache_key] = text
    return text


def _parse_tz_pdf(path: Path) -> list[TZItem]:
    items = _parse_pdf_with_pdfplumber(path)
    if items:
        return items

    text = _extract_pdf_full_text(path)
    items = _parse_price_request_text(text)
    if items:
        return items

    items = _parse_numbered_list_text(text)
    if items:
        return items

    items = _parse_tz_tables(_tables_from_text(text))
    if items:
        return items

    if _pdf_page_is_image_only(path) and not _tesseract_available():
        raise ValueError(
            "PDF загружен как изображение (скан) без текстового слоя. "
            "Для распознавания нужен Tesseract OCR "
            "(Docker: tesseract-ocr tesseract-ocr-rus; macOS: brew install tesseract tesseract-lang)."
        )

    if _looks_like_procurement_cover_letter(text) or (
        _pdf_page_is_image_only(path) and _ocr_text_looks_like_official_letter(text)
    ):
        raise ValueError(
            "PDF содержит сопроводительное письмо (запрос КП) без списка позиций. "
            "Загрузите приложение с перечнем оборудования или файл со списком позиций."
        )

    raise ValueError(
        "Не удалось извлечь позиции из PDF. "
        "Поддерживаются таблицы с колонкой «Наименование» и нумерованные списки (1. …, 2. …). "
        "Для сканов установите Tesseract и включите TZ_PDF_OCR_ENABLED=true."
    )


def _extract_pdf_document_text(path: Path) -> str:
    return _extract_pdf_full_text(path)


def _parse_pdf_with_pdfplumber(path: Path) -> list[TZItem]:
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber не установлен — пропускаю извлечение таблиц из PDF")
        return []

    tables: list[list[list[str]]] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                cleaned = [[_cell_str(cell) for cell in row] for row in table if row]
                if cleaned:
                    tables.append(cleaned)

            page_text = page.extract_text() or ""
            if page_text:
                tables.extend(_tables_from_text(page_text))

    return _parse_tz_tables(tables)


def _extract_pdf_text(path: Path) -> str:
    try:
        import fitz
    except ImportError:
        logger.warning("pymupdf не установлен — пропускаю извлечение текста из PDF")
        return ""

    chunks: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            chunks.append(page.get_text("text"))
    return "\n".join(chunks)


def _configure_tesseract() -> None:
    if not TESSERACT_CMD:
        return
    import pytesseract

    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


def _tesseract_available() -> bool:
    try:
        import pytesseract

        _configure_tesseract()
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _preprocess_ocr_image(image):
    from PIL import ImageOps

    gray = ImageOps.grayscale(image)
    return ImageOps.autocontrast(gray)


def _ocr_pdf(path: Path) -> str:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise ValueError(
            "Для OCR PDF установите зависимости: pymupdf, pytesseract, Pillow"
        ) from exc

    _configure_tesseract()
    if not _tesseract_available():
        raise ValueError(
            "Tesseract OCR не найден. Установите: brew install tesseract tesseract-lang "
            "(macOS) или apt install tesseract-ocr tesseract-ocr-rus (Linux)."
        )

    config = f"--psm {TZ_OCR_PSM}" if TZ_OCR_PSM > 0 else ""
    texts: list[str] = []
    matrix = fitz.Matrix(TZ_OCR_SCALE, TZ_OCR_SCALE)
    with fitz.open(path) as doc:
        for page in doc:
            pix = page.get_pixmap(matrix=matrix)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = _preprocess_ocr_image(img)
            texts.append(
                pytesseract.image_to_string(
                    img,
                    lang=TZ_OCR_LANG,
                    config=config,
                )
            )

    return "\n".join(texts)


def _tables_from_text(text: str) -> list[list[list[str]]]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    tables: list[list[list[str]]] = []
    idx = 0

    while idx < len(lines):
        if "наимен" not in lines[idx].lower():
            idx += 1
            continue

        header = _split_table_line(lines[idx])
        if not _is_tz_table_header(header):
            idx += 1
            continue
        col = _build_column_map(header)
        rows = [header]
        idx += 1

        while idx < len(lines):
            line = lines[idx]
            if "наимен" in line.lower() and rows:
                break
            parts = _split_table_line(line)
            has_number = col.has_number if col else True
            if _looks_like_data_row(parts, has_number=has_number):
                rows.append(parts)
                idx += 1
                continue
            if rows and len(rows) > 1:
                break
            idx += 1

        if len(rows) > 1:
            tables.append(rows)

    return tables


def _split_table_line(line: str) -> list[str]:
    if "\t" in line:
        return [part.strip() for part in line.split("\t") if part.strip()]
    parts = re.split(r"\s{2,}", line)
    if len(parts) >= 4:
        return [part.strip() for part in parts]
    return [part.strip() for part in line.split() if part.strip()]


def _looks_like_data_row(parts: list[str], *, has_number: bool = True) -> bool:
    if len(parts) < 2:
        return False
    if not has_number:
        return bool(parts[0].strip())
    number = parts[0].rstrip(".")
    return number.isdigit() and bool(parts[1])
